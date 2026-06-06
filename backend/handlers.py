import os
import zipfile
import tempfile
import shutil
import requests
from xml.etree import ElementTree as ET
import fitz  # PyMuPDF
import docx
from fpdf import FPDF
from converter import convert_nepali_text

# ==========================================
# DEVANAGARI FONT DOWNLOAD SYSTEM FOR PDF
# ==========================================

FONT_URL = "https://raw.githubusercontent.com/googlefonts/noto-fonts/main/hinted/ttf/NotoSansDevanagari/NotoSansDevanagari-Regular.ttf"
FONT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "NotoSansDevanagari-Regular.ttf")

def ensure_devanagari_font():
    """
    Dynamically downloads Noto Sans Devanagari from Google Fonts if it's not cached locally.
    This font is required to render Unicode Devanagari characters in FPDF2.
    """
    if not os.path.exists(FONT_PATH):
        try:
            print("Downloading Noto Sans Devanagari font...")
            r = requests.get(FONT_URL, timeout=15)
            if r.status_code == 200:
                with open(FONT_PATH, "wb") as f:
                    f.write(r.content)
                print("Font NotoSansDevanagari downloaded successfully.")
            else:
                print(f"Failed to download font: status code {r.status_code}")
        except Exception as e:
            print(f"Error downloading font: {e}")
    return FONT_PATH if os.path.exists(FONT_PATH) else None


# ==========================================
# FORMAT HANDLERS
# ==========================================

def convert_txt(input_path: str, output_path: str, direction: str):
    """
    Converts a plain text file. Supports robust fallback encoding for legacy ANSI text files.
    """
    try:
        with open(input_path, "r", encoding="utf-8") as f:
            content = f.read()
    except UnicodeDecodeError:
        # Fallback to standard ANSI (cp1252/latin-1) often used for legacy Preeti text
        with open(input_path, "r", encoding="cp1252") as f:
            content = f.read()
            
    converted_content = convert_nepali_text(content, direction)
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(converted_content)


def convert_docx(input_path: str, output_path: str, direction: str):
    """
    Converts a DOCX Word document.
    Iterates through runs in paragraphs, tables, headers, and footers, and maps text.
    Preserves styles, bold/italic, alignments, and changes font names to display correctly.
    """
    doc = docx.Document(input_path)
    
    # We choose Mangal for Unicode (natively supported on Windows) or Preeti for Preeti
    font_name = "Mangal" if direction == "preeti_to_unicode" else "Preeti"
    
    def process_paragraph(p):
        for run in p.runs:
            if run.text and run.text.strip():
                # Convert run text
                run.text = convert_nepali_text(run.text, direction)
                run.font.name = font_name
                
                # Directly update openxml font tags for proper rendering in complex scripts (w:cs)
                rPr = run._r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(docx.oxml.ns.qn('w:ascii'), font_name)
                rFonts.set(docx.oxml.ns.qn('w:hAnsi'), font_name)
                rFonts.set(docx.oxml.ns.qn('w:cs'), font_name)

    # 1. Process main paragraphs
    for p in doc.paragraphs:
        process_paragraph(p)
        
    # 2. Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
                    
    # 3. Process headers and footers
    for section in doc.sections:
        if section.header and not section.header.is_linked_to_previous:
            for p in section.header.paragraphs:
                process_paragraph(p)
        if section.footer and not section.footer.is_linked_to_previous:
            for p in section.footer.paragraphs:
                process_paragraph(p)
                
    doc.save(output_path)


def convert_odt(input_path: str, output_path: str, direction: str):
    """
    Converts ODT documents.
    Unzips the file, parses content.xml, maps text and tail elements, and re-zips.
    This lightweight XML parsing guarantees 100% portable ODT styling preservation.
    """
    tmp_dir = tempfile.mkdtemp()
    try:
        # Extract ODT structure
        with zipfile.ZipFile(input_path, 'r') as zip_ref:
            zip_ref.extractall(tmp_dir)
            
        content_xml_path = os.path.join(tmp_dir, 'content.xml')
        if os.path.exists(content_xml_path):
            # Register namespaces to prevent ElementTree from stripping namespace prefixes
            namespaces = {
                'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
                'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
                'style': 'urn:oasis:names:tc:opendocument:xmlns:style:1.0',
                'table': 'urn:oasis:names:tc:opendocument:xmlns:table:1.0',
                'fo': 'urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0',
                'svg': 'urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0'
            }
            for prefix, uri in namespaces.items():
                ET.register_namespace(prefix, uri)
                
            tree = ET.parse(content_xml_path)
            root = tree.getroot()
            
            # Map all text nodes
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    elem.text = convert_nepali_text(elem.text, direction)
                if elem.tail and elem.tail.strip():
                    elem.tail = convert_nepali_text(elem.tail, direction)
                    
            tree.write(content_xml_path, encoding='utf-8', xml_declaration=True)
            
        # Re-pack ODT file
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root_dir, dirs, files in os.walk(tmp_dir):
                for file in files:
                    full_path = os.path.join(root_dir, file)
                    rel_path = os.path.relpath(full_path, tmp_dir)
                    zip_out.write(full_path, rel_path)
    finally:
        shutil.rmtree(tmp_dir)


def convert_pdf(input_path: str, output_path: str, direction: str):
    """
    Converts a PDF file.
    Extracts text layout using PyMuPDF, maps it, and regenerates a new PDF using FPDF2.
    Renders a clear red header disclaimer at the top of the reconstructed PDF.
    """
    # 1. Extract text using PyMuPDF
    doc = fitz.open(input_path)
    extracted_pages = []
    for page in doc:
        extracted_pages.append(page.get_text("text"))
    doc.close()
    
    # 2. Setup custom FPDF subclass to output disclaimer
    class ConvertedPDF(FPDF):
        def header(self):
            self.set_text_color(220, 53, 69) # Bootstrap Danger Red
            self.set_font('helvetica', 'BI', 9)
            disclaimer = "DISCLAIMER: Converted between Unicode & Preeti. Layout shifts or formatting differences may occur."
            self.cell(0, 8, disclaimer, 0, 1, 'C')
            self.ln(3)
            
    pdf = ConvertedPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Configure font based on output
    if direction == "preeti_to_unicode":
        # Unicode Devanagari requires a TrueType Font
        font_file = ensure_devanagari_font()
        if font_file:
            pdf.add_font("NotoSansDevanagari", style="", fname=font_file)
            pdf.set_font("NotoSansDevanagari", size=11)
        else:
            pdf.set_font("helvetica", size=11)
    else:
        # Preeti font is ASCII character-based, rendering natively in Helvetica
        pdf.set_font("helvetica", size=11)
        
    pdf.set_text_color(33, 37, 41) # Dark gray text
    
    # 3. Add mapped text to PDF pages
    for page_text in extracted_pages:
        pdf.add_page()
        converted_text = convert_nepali_text(page_text, direction)
        
        # Write lines to PDF safely
        for line in converted_text.split('\n'):
            # Convert to latin-1 to avoid encoding issues in standard fpdf latin character mappings
            clean_line = line.strip()
            if clean_line:
                # FPDF2 cell multi-line handling
                pdf.multi_cell(0, 6, clean_line)
                pdf.ln(1)
                
    pdf.output(output_path)
